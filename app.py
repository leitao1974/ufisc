import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from io import BytesIO

# Configuração da Página
st.set_page_config(page_title="Fiscalização Digital v2", layout="wide", page_icon="⚖️")

# --- LISTA OFICIAL DE TIPOLOGIAS REN ---
TIPOLOGIAS_REN = [
    "--- Selecione uma Tipologia ---",
    "Estrutura de proteção e recarga de aquíferos (Zonas de infiltração máxima)",
    "Áreas estratégicas de proteção e recarga de aquíferos",
    "Zonas adjacentes",
    "Zonas inundáveis (Ameaçadas pelas cheias)",
    "Albufeiras (Faixas de proteção)",
    "Cursos de água (Leitos e margens)",
    "Cabeceiras de linhas de água",
    "Arribas e respetivas faixas de proteção",
    "Praias e dunas",
    "Escarpas e respetivas faixas de proteção",
    "Áreas de elevada perigosidade de incêndio florestal",
    "Áreas de instabilidade de vertentes (Movimentos de massa)",
    "Estuários e Lagunas",
    "Zonas húmidas"
]

def extrair_texto_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    texto = ""
    for page in reader.pages:
        texto += page.extract_text()
    return texto

def criar_word(texto_parecer):
    doc = Document()
    doc.add_heading('Parecer Técnico-Jurídico de Fiscalização', 0)
    for linha in texto_parecer.split('\n'):
        if linha.strip():
            p = doc.add_paragraph()
            limpa_linha = linha.replace('**', '').replace('###', '').strip()
            run = p.add_run(limpa_linha)
            if '**' in linha:
                run.bold = True
    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return target

## --- SIDEBAR ---
with st.sidebar:
    st.title("⚙️ Configurações")
    api_key = st.text_input("Introduza a sua Google API Key:", type="password")
    
    if api_key:
        try:
            genai.configure(api_key=api_key)
            available_models = [m.name.replace('models/', '') for m in genai.list_models() 
                                if 'generateContent' in m.supported_generation_methods]
            model_choice = st.selectbox("Selecione o Modelo Disponível:", available_models)
        except Exception as e:
            st.error(f"Erro ao validar chave: {e}")
            model_choice = "gemini-1.5-pro"
    else:
        st.warning("Insira a API Key para listar os modelos.")
        model_choice = "gemini-1.5-pro"

    st.divider()
    uploaded_file = st.file_uploader("Upload do Auto de Notícia (PDF)", type="pdf")

## --- CORPO DA APP ---
st.header("⚖️ Sistema de Apoio à Fiscalização")

st.subheader("Enquadramentos Jurídicos")
c1, c2, c3 = st.columns(3)

with c1:
    check_ren = st.checkbox("REN (DL 166/2008 + DL 123/2024)")
    tipologia_selecionada = ""
    if check_ren:
        tipologia_selecionada = st.selectbox("Selecione a Tipologia REN:", TIPOLOGIAS_REN)
    
    check_ran = st.checkbox("RAN (DL 73/2009 + DL 199/2015)")

with c2:
    check_natura = st.checkbox("Rede Natura 2000 (DL 49/2005)")
    check_patrimonio = st.checkbox("Património (Lei 107/2001)")

with c3:
    check_rjue = st.checkbox("Ordenamento/Urbanismo (RJUE + DL 10/2024)")
    check_coimas = st.checkbox("Coimas (Lei 50/2006 + DL 87/2024)")

if st.button("🚀 Gerar Parecer Jurídico"):
    if not api_key or not uploaded_file:
        st.error("⚠️ Falta a API Key ou o ficheiro PDF.")
    elif check_ren and (tipologia_selecionada == TIPOLOGIAS_REN[0]):
        st.warning("⚠️ Selecione uma tipologia REN válida.")
    else:
        with st.spinner(f"A analisar regimes com {model_choice}..."):
            try:
                texto_auto = extrair_texto_pdf(uploaded_file)
                model = genai.GenerativeModel(model_name=f"models/{model_choice}")

                diretrizes = []
                
                # REFORÇO DA ANÁLISE RAN
                if check_ran:
                    diretrizes.append(
                        "- RAN (Reserva Agrícola Nacional): Aplicar DL 73/2009 (redação DL 199/2015).\n"
                        "  1. Identificar se a utilização é não agrícola (Art. 21.º).\n"
                        "  2. Transcrever e analisar as alíneas do Art. 21.º ou 22.º aplicáveis.\n"
                        "  3. Verificar requisitos cumulativos para legalização: inexistência de alternativa fora da RAN e razões de interesse público."
                    )

                if check_ren: 
                    diretrizes.append(f"- REN: DL 166/2008 e DL 123/2024. Tipologia: '{tipologia_selecionada}'. Fundamentar com Anexo I e Anexo II.")
                
                if check_natura: 
                    diretrizes.append("- REDE NATURA 2000: DL 49/2005. Analisar Art. 9.º, n.º 2 (alíneas a-l) e Secção III (Espécies).")
                
                if check_rjue: diretrizes.append("- Urbanismo: RJUE e DL 10/2024. Verificar Nulidade do Art. 68.º.")
                
                if check_coimas: diretrizes.append("- Coimas: Lei 50/2006 e DL 87/2024. Identificar gravidade e limites das coimas.")

                prompt_final = f"""
                Age como um Jurista Sénior especialista em Ordenamento em Portugal (PT-PT).
                Analisa o AUTO DE NOTÍCIA com base nestas diretrizes:
                {chr(10).join(diretrizes)}

                REGRAS CRÍTICAS:
                1. RAN: Se selecionada, deves obrigatoriamente verificar se a ação constitui uma utilização não agrícola proibida sem parecer da entidade regional (Art. 21.º). Concluir se é suscetível de regularização excecional.
                2. REN: Justificar a violação das FUNÇÕES do ANEXO I para a tipologia '{tipologia_selecionada}'.
                3. REDE NATURA: Citar a alínea específica do Art. 9.º, n.º 2 do DL 49/2005.
                4. NULIDADES: Invocar o Art. 68.º do RJUE se houver omissão de pareceres da RAN, REN ou ICNF.
                5. LEGALIZAÇÃO: Analisar de forma taxativa: "Legalizável" ou "Insuscetível de Legalização".

                TEXTO DO AUTO:
                {texto_auto}

                ESTRUTURA OBRIGATÓRIA:
                1. **OBJECTIVO**
                2. **DESCRIÇÃO TÉCNICA E AUDITORIA**
                3. **FUNDAMENTAÇÃO JURÍDICA E TRANSGRESSÕES** (Analisar cada regime selecionado separadamente)
                4. **ANÁLISE JURÍDICA DE VIABILIDADE DE LEGALIZAÇÃO**
                5. **QUADRO SANCIONATÓRIO E NULIDADES**
                6. **PARECER FINAL E MEDIDAS DE REPOSIÇÃO**
                """

                response = model.generate_content(prompt_final)
                parecer_texto = response.text
                
                st.markdown("---")
                st.subheader("📄 Parecer Jurídico Gerado")
                st.markdown(parecer_texto)
                
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    st.download_button("Baixar em Markdown (.md)", parecer_texto, file_name="parecer.md")
                with col_d2:
                    word_file = criar_word(parecer_texto)
                    st.download_button("Baixar em Word (.docx)", word_file, file_name="parecer_fiscalizacao.docx")

            except Exception as e:
                st.error(f"Erro: {e}")
